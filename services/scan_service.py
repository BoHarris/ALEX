import html as html_lib
import io
import json
import logging
import mimetypes
import os
import re
import uuid
import warnings
import xml.etree.ElementTree as ET
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterator, Optional

import joblib
import numpy as np
import pandas as pd
with warnings.catch_warnings():
    warnings.filterwarnings("ignore", message="PyPDF2 is deprecated.*", category=DeprecationWarning)
    from PyPDF2 import PdfReader
from docx import Document
from sqlalchemy.exc import IntegrityError

from database.database import SessionLocal
from database.models.scan_results import ScanResult
from services.retention_service import build_retention_expiration, resolve_retention_days
from services.audit_service import record_audit_event
from utils.constants import SUPPORTED_EXTENSIONS
from utils.pii_taxonomy import (
    PII_ADDRESS,
    PII_DATE_OF_BIRTH,
    PII_EMAIL,
    PII_IP_ADDRESS,
    PII_NAME,
    PII_PHONE,
    PII_SENSITIVE_DATA,
    PII_SSN,
    get_display_name,
    get_policy_categories,
    map_legacy_label_to_taxonomy,
)
from utils.redaction import scan_and_redact_column_with_details

logger = logging.getLogger(__name__)

BASE_DIR = Path(__file__).resolve().parent.parent
MODEL_PATH = BASE_DIR / "models" / "xgboost_model.pkl"
_xgb_model = None

EMAIL_RE = re.compile(r"[^@]+@[^@]+\.[^@]+")
PHONE_RE = re.compile(r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}")
SSN_RE = re.compile(r"\d{3}-\d{2}-\d{4}")
IP_RE = re.compile(r"(?:\d{1,3}\.){3}\d{1,3}")
STREET_SUFFIXES = {"st", "street", "ave", "road", "rd", "blvd", "ln", "lane"}
MAX_FEATURE_SAMPLE_VALUES = max(3, int(os.getenv("FEATURE_SAMPLE_VALUES", "15")))
SCAN_CHUNK_ROWS = max(100, int(os.getenv("SCAN_CHUNK_ROWS", "1000")))
MAX_SCAN_ROWS = max(1, int(os.getenv("MAX_SCAN_ROWS", "50000")))
MAX_SCAN_CELLS = max(1, int(os.getenv("MAX_SCAN_CELLS", "500000")))
MAX_SCAN_COLUMNS = max(1, int(os.getenv("MAX_SCAN_COLUMNS", "200")))
XML_TAG_INVALID_CHARS_RE = re.compile(r"[^A-Za-z0-9_.-]+")


@dataclass(frozen=True)
class ScanContext:
    user_id: int
    company_id: Optional[int] = None
    tier: Optional[str] = None


@dataclass(frozen=True)
class ScanPipelineResult:
    filename: str
    pii_columns: list[str]
    redacted_file: str
    risk_score: int
    redacted_count: int
    total_values: int
    redacted_type_counts: dict[str, int]
    detection_results: list[dict[str, object]]
    scan_id: int


@dataclass(frozen=True)
class ScanLimits:
    max_rows: int = MAX_SCAN_ROWS
    max_cells: int = MAX_SCAN_CELLS
    max_columns: int = MAX_SCAN_COLUMNS
    chunk_rows: int = SCAN_CHUNK_ROWS


@dataclass(frozen=True)
class DetectionResult:
    column: str
    detected_type: str
    display_name: str
    policy_categories: list[str]
    confidence_score: float
    signals: list[str]


class ScanLimitError(ValueError):
    def __init__(self, detail: str, *, error_code: str = "FILE_TOO_LARGE", status_code: int = 413):
        super().__init__(detail)
        self.detail = detail
        self.error_code = error_code
        self.status_code = status_code


def _load_xgboost_model(model_path: Path = MODEL_PATH):
    try:
        return joblib.load(model_path)
    except Exception as exc:
        raise RuntimeError(
            "StartupError: XGBoost model could not be loaded. "
            "Verify model file exists and is compatible with current environment."
        ) from exc


def initialize_scan_model(model_path: Path = MODEL_PATH):
    global _xgb_model
    if _xgb_model is None:
        _xgb_model = _load_xgboost_model(model_path)
    return _xgb_model


def set_scan_model(model) -> None:
    global _xgb_model
    _xgb_model = model


def get_scan_model():
    if _xgb_model is None:
        raise RuntimeError(
            "Scan model is not initialized. Startup validation must load the XGBoost model before scanning."
        )
    return _xgb_model


def _contains_dob_pattern(values: list[str]) -> bool:
    dob_regex = re.compile(r"\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b")
    return any(dob_regex.search(str(value)) for value in values) if values else False


def _contains_gender_term(values: list[str]) -> bool:
    gender_terms = {"male", "female", "man", "woman", "boy", "girl"}
    return any(any(term in str(value).lower() for term in gender_terms) for value in values) if values else False


def _contains_street_suffix(values: list[str]) -> bool:
    return any(any(suffix in str(value).lower() for suffix in STREET_SUFFIXES) for value in values) if values else False


def _contains_city_name(values: list[str]) -> bool:
    city_names = {"new york", "los angeles", "chicago", "houston", "phoenix"}
    return any(any(city in str(value).lower() for city in city_names) for value in values) if values else False


def _contains_known_name(values: list[str]) -> bool:
    known_names = {"john", "jane", "smith", "doe"}
    return any(any(name in str(value).lower() for name in known_names) for value in values) if values else False


def _contains_zip_code_pattern(values: list[str]) -> bool:
    zip_regex = re.compile(r"\b\d{5}(?:-\d{4})?\b")
    return any(zip_regex.search(str(value)) for value in values) if values else False


def _contains_phone_pattern(values: list[str]) -> bool:
    return any(PHONE_RE.search(str(value)) for value in values) if values else False


def _contains_email_pattern(values: list[str]) -> bool:
    return any(EMAIL_RE.search(str(value)) for value in values) if values else False


def _contains_ssn_pattern(values: list[str]) -> bool:
    return any(SSN_RE.search(str(value)) for value in values) if values else False


def _contains_ip_pattern(values: list[str]) -> bool:
    return any(IP_RE.search(str(value)) for value in values) if values else False


def _column_sample_values(series: pd.Series) -> list[str]:
    values = series.dropna().astype(str)
    if values.empty:
        return []
    if len(values) <= MAX_FEATURE_SAMPLE_VALUES:
        return values.tolist()
    indices = np.linspace(0, len(values) - 1, num=MAX_FEATURE_SAMPLE_VALUES, dtype=int)
    return values.iloc[indices].tolist()


def _normalize_column_tokens(column_name: str) -> set[str]:
    normalized = re.sub(r"[^a-z0-9]+", " ", str(column_name).lower())
    return {token for token in normalized.split() if token}


def _column_has_keyword(column_name: str, *keywords: str) -> bool:
    column_tokens = _normalize_column_tokens(column_name)
    normalized_text = re.sub(r"[^a-z0-9]+", " ", str(column_name).lower())
    for keyword in keywords:
        if " " in keyword:
            if keyword in normalized_text:
                return True
        elif keyword in column_tokens:
            return True
    return False


def _candidate_detection(
    *,
    column_name: str,
    values: list[str],
    ml_detected: bool,
    detected_type: str,
    display_name: str,
    keyword_signal: tuple[str, tuple[str, ...], float] | None = None,
    pattern_signal: tuple[str, bool, float] | None = None,
    format_signal: tuple[str, bool, float] | None = None,
    context_signal: tuple[str, bool, float] | None = None,
) -> DetectionResult | None:
    signals: list[str] = []
    score = 0.0

    if keyword_signal and _column_has_keyword(column_name, *keyword_signal[1]):
        signals.append(keyword_signal[0])
        score += keyword_signal[2]
    if pattern_signal and pattern_signal[1]:
        signals.append(pattern_signal[0])
        score += pattern_signal[2]
    if format_signal and format_signal[1]:
        signals.append(format_signal[0])
        score += format_signal[2]
    if context_signal and context_signal[1]:
        signals.append(context_signal[0])
        score += context_signal[2]
    if ml_detected and signals:
        signals.append("ml_model_prediction")
        score += 0.18

    if not signals:
        return None

    return DetectionResult(
        column=column_name,
        detected_type=detected_type,
        display_name=display_name,
        policy_categories=get_policy_categories(detected_type),
        confidence_score=round(min(score, 0.99), 2),
        signals=signals,
    )


def _build_detection_result(*, column_name: str, values: list[str], ml_detected: bool) -> DetectionResult | None:
    candidates = [
        _candidate_detection(
            column_name=column_name,
            values=values,
            ml_detected=ml_detected,
            detected_type=PII_EMAIL,
            display_name=get_display_name(PII_EMAIL),
            keyword_signal=("column_name_keyword_email", ("email", "e mail"), 0.24),
            pattern_signal=("pattern_match_email_regex", _contains_email_pattern(values), 0.44),
            format_signal=("value_format_match_email", _contains_email_pattern(values), 0.12),
        ),
        _candidate_detection(
            column_name=column_name,
            values=values,
            ml_detected=ml_detected,
            detected_type=PII_PHONE,
            display_name=get_display_name(PII_PHONE),
            keyword_signal=("column_name_keyword_phone", ("phone", "telephone", "mobile", "cell", "tel"), 0.24),
            pattern_signal=("pattern_match_phone_regex", _contains_phone_pattern(values), 0.42),
            format_signal=("value_format_match_phone", _contains_phone_pattern(values), 0.14),
        ),
        _candidate_detection(
            column_name=column_name,
            values=values,
            ml_detected=ml_detected,
            detected_type=PII_SSN,
            display_name=get_display_name(PII_SSN),
            keyword_signal=("column_name_keyword_ssn", ("ssn", "social security"), 0.24),
            pattern_signal=("pattern_match_ssn_regex", _contains_ssn_pattern(values), 0.5),
        ),
        _candidate_detection(
            column_name=column_name,
            values=values,
            ml_detected=ml_detected,
            detected_type=PII_IP_ADDRESS,
            display_name=get_display_name(PII_IP_ADDRESS),
            keyword_signal=("column_name_keyword_ip", ("ip", "ip address", "ipv4", "ipv6"), 0.24),
            pattern_signal=("pattern_match_ip_regex", _contains_ip_pattern(values), 0.5),
        ),
        _candidate_detection(
            column_name=column_name,
            values=values,
            ml_detected=ml_detected,
            detected_type=PII_ADDRESS,
            display_name=get_display_name(PII_ADDRESS),
            keyword_signal=("column_name_keyword_address", ("address", "street", "city", "zip", "postal"), 0.22),
            format_signal=("value_format_match_address", _contains_street_suffix(values), 0.26),
            context_signal=("context_match_address_geography", _contains_city_name(values) or _contains_zip_code_pattern(values), 0.16),
        ),
        _candidate_detection(
            column_name=column_name,
            values=values,
            ml_detected=ml_detected,
            detected_type=PII_NAME,
            display_name=get_display_name(PII_NAME),
            keyword_signal=("column_name_keyword_name", ("name", "first name", "last name", "full name"), 0.22),
            format_signal=("value_format_match_name", _contains_known_name(values), 0.24),
        ),
        _candidate_detection(
            column_name=column_name,
            values=values,
            ml_detected=ml_detected,
            detected_type=PII_DATE_OF_BIRTH,
            display_name=get_display_name(PII_DATE_OF_BIRTH),
            keyword_signal=("column_name_keyword_dob", ("dob", "birth", "date of birth"), 0.22),
            pattern_signal=("pattern_match_dob_regex", _contains_dob_pattern(values), 0.36),
            context_signal=("context_match_age_sensitive_column", _column_has_keyword(column_name, "dob", "birth"), 0.12),
        ),
    ]

    scored_candidates = [candidate for candidate in candidates if candidate is not None]
    if scored_candidates:
        return max(scored_candidates, key=lambda candidate: (candidate.confidence_score, candidate.display_name))
    if ml_detected:
        return DetectionResult(
            column=column_name,
            detected_type=PII_SENSITIVE_DATA,
            display_name=get_display_name(PII_SENSITIVE_DATA),
            policy_categories=get_policy_categories(PII_SENSITIVE_DATA),
            confidence_score=0.18,
            signals=["ml_model_prediction"],
        )
    return None


def build_scan_result_metadata(
    *,
    redacted_type_counts: dict[str, int],
    detection_results: list[dict[str, object]],
) -> dict[str, object]:
    normalized_counts: dict[str, int] = {}
    for label, count in redacted_type_counts.items():
        taxonomy_code = map_legacy_label_to_taxonomy(label)
        normalized_counts[taxonomy_code] = normalized_counts.get(taxonomy_code, 0) + int(count)
    return {
        "counts": normalized_counts,
        "detections": list(detection_results),
    }


def parse_scan_result_metadata(raw_value: str | None) -> tuple[dict[str, int], list[dict[str, object]]]:
    if not raw_value:
        return {}, []
    try:
        data = json.loads(raw_value)
    except (TypeError, json.JSONDecodeError):
        return {}, []
    if not isinstance(data, dict):
        return {}, []

    raw_counts = data.get("counts", data)
    counts: dict[str, int] = {}
    if isinstance(raw_counts, dict):
        for key, value in raw_counts.items():
            taxonomy_code = map_legacy_label_to_taxonomy(key)
            if not taxonomy_code:
                continue
            try:
                counts[taxonomy_code] = counts.get(taxonomy_code, 0) + int(value)
            except (TypeError, ValueError):
                continue

    detections: list[dict[str, object]] = []
    raw_detections = data.get("detections", [])
    if isinstance(raw_detections, list):
        for item in raw_detections:
            if not isinstance(item, dict):
                continue
            column = str(item.get("column", "")).strip()
            detected_type = str(item.get("detected_type") or item.get("detected_as") or "").strip()
            display_name = str(item.get("display_name", "")).strip() or get_display_name(detected_type)
            policy_categories = item.get("policy_categories", get_policy_categories(detected_type))
            signals = item.get("signals", [])
            if not column or not detected_type or not isinstance(signals, list):
                continue
            try:
                confidence_score = round(float(item.get("confidence_score", 0.0)), 2)
            except (TypeError, ValueError):
                continue
            if not isinstance(policy_categories, list):
                policy_categories = get_policy_categories(detected_type)
            detections.append(
                {
                    "column": column,
                    "detected_type": detected_type,
                    "detected_as": detected_type,
                    "display_name": display_name,
                    "policy_categories": [str(policy) for policy in policy_categories if str(policy).strip()],
                    "confidence_score": max(0.0, min(confidence_score, 1.0)),
                    "signals": [str(signal) for signal in signals if str(signal).strip()],
                }
            )
    return counts, detections


def _raise_scan_limit(detail: str) -> None:
    raise ScanLimitError(detail=detail)


def _enforce_tabular_limits(*, row_count: int, column_count: int, limits: ScanLimits) -> None:
    if column_count > limits.max_columns:
        _raise_scan_limit("Uploaded file exceeds allowed processing limits")
    if row_count > limits.max_rows:
        _raise_scan_limit("Uploaded file exceeds allowed processing limits")
    if row_count * max(column_count, 1) > limits.max_cells:
        _raise_scan_limit("Uploaded file exceeds allowed processing limits")


def _read_source_bytes(*, file_bytes: bytes | None, source_path: str | None) -> bytes:
    if file_bytes is not None:
        return file_bytes
    if not source_path:
        raise ValueError("Scan source is required.")
    with open(source_path, "rb") as handle:
        return handle.read()


def _parse_to_dataframe(
    *,
    file_bytes: bytes | None,
    source_path: str | None,
    filename: str,
    ext: str,
    limits: ScanLimits,
) -> pd.DataFrame:
    source_bytes = _read_source_bytes(file_bytes=file_bytes, source_path=source_path)
    file_format, _ = mimetypes.guess_type(filename)
    file_format = file_format or ""

    if ext == ".pdf":
        reader = PdfReader(io.BytesIO(source_bytes))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise ValueError("PDF is encrypted or requires a password.") from exc
        texts = [page.extract_text() or "" for page in reader.pages]
        parsed = pd.DataFrame({"text": texts})
    elif ext == ".csv" or "csv" in file_format:
        parsed = pd.read_csv(io.BytesIO(source_bytes))
    elif ext == ".tsv":
        parsed = pd.read_csv(io.BytesIO(source_bytes), sep="\t")
    elif ext in {".xls", ".xlsx"} or "excel" in file_format:
        parsed = pd.read_excel(io.BytesIO(source_bytes))
    elif ext == ".json" or "json" in file_format:
        data = json.load(io.BytesIO(source_bytes))
        parsed = pd.json_normalize(data)
    elif ext in {".txt", ".log"} or "plain" in file_format:
        lines = io.BytesIO(source_bytes).read().decode("utf-8", errors="ignore").splitlines()
        parsed = pd.DataFrame({"text": lines})
    elif ext == ".xml" or "xml" in file_format:
        tree = ET.parse(io.BytesIO(source_bytes))
        root = tree.getroot()
        parsed = pd.DataFrame([{child.tag: child.text for child in elem} for elem in root])
    elif ext == ".docx":
        doc = Document(io.BytesIO(source_bytes))
        chunks: list[str] = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                chunks.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        chunks.append(cell.text)
        parsed = pd.DataFrame({"text": chunks})
    elif ext == ".html":
        try:
            tables = pd.read_html(io.BytesIO(source_bytes))
            parsed = tables[0] if tables else pd.DataFrame({"text": []})
        except Exception:
            raw = io.BytesIO(source_bytes).read().decode("utf-8", errors="ignore")
            raw = re.sub(r"(?is)<(script|style).*?>.*?</\1>", "", raw)
            text = re.sub(r"(?s)<.*?>", "\n", raw)
            lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
            parsed = pd.DataFrame({"text": lines})
    else:
        raise ValueError(f"Unsupported format: {file_format} / {ext}")

    df = parsed if isinstance(parsed, pd.DataFrame) else pd.DataFrame({"text": parsed})
    if df.empty:
        raise ValueError("File is empty or not supported.")
    _enforce_tabular_limits(row_count=len(df), column_count=len(df.columns), limits=limits)
    return df


def _iter_csv_like_chunks(*, source_path: str, ext: str, chunk_rows: int) -> Iterator[pd.DataFrame]:
    separator = "\t" if ext == ".tsv" else ","
    with open(source_path, "r", encoding="utf-8", errors="ignore", newline="") as handle:
        for chunk in pd.read_csv(handle, chunksize=chunk_rows, sep=separator):
            yield chunk


def _iter_text_chunks(*, source_path: str, chunk_rows: int) -> Iterator[pd.DataFrame]:
    batch: list[str] = []
    with open(source_path, "r", encoding="utf-8", errors="ignore") as handle:
        for line in handle:
            batch.append(line.rstrip("\n"))
            if len(batch) >= chunk_rows:
                yield pd.DataFrame({"text": batch})
                batch = []
    if batch:
        yield pd.DataFrame({"text": batch})


def _build_feature_dataframe_from_samples(features: list[str], parsed_values: list[list[str]]) -> pd.DataFrame:
    x = pd.DataFrame(features, columns=["column"])
    x["length"] = x["column"].apply(len)
    x["num_underscores"] = x["column"].apply(lambda col: col.count("_"))
    x["num_digits"] = x["column"].apply(lambda col: sum(c.isdigit() for c in col))
    x["has_at"] = x["column"].apply(lambda col: int("@" in col))
    x["has_email_keyword"] = x["column"].apply(lambda col: int("email" in col.lower()))
    x["parsed_values"] = parsed_values

    def pct_match(regex: re.Pattern, values: list[str]) -> float:
        return float(np.mean([bool(regex.search(str(v))) for v in values])) if values else 0.0

    x["pct_email_like"] = x["parsed_values"].apply(lambda vals: pct_match(EMAIL_RE, vals))
    x["pct_phone_like"] = x["parsed_values"].apply(lambda vals: pct_match(PHONE_RE, vals))
    x["pct_ssn_like"] = x["parsed_values"].apply(lambda vals: pct_match(SSN_RE, vals))
    x["pct_ip_like"] = x["parsed_values"].apply(lambda vals: pct_match(IP_RE, vals))
    x["avg_digits_per_val"] = x["parsed_values"].apply(
        lambda vals: np.mean([sum(c.isdigit() for c in str(item)) for item in vals]) if vals else 0
    )
    x["avg_val_len"] = x["parsed_values"].apply(
        lambda vals: np.mean([len(str(item)) for item in vals]) if vals else 0
    )
    x["has_dob_pattern"] = x["parsed_values"].apply(_contains_dob_pattern)
    x["has_gender_term"] = x["parsed_values"].apply(_contains_gender_term)
    x["has_street_suffix"] = x["parsed_values"].apply(_contains_street_suffix)
    x["has_city_name"] = x["parsed_values"].apply(_contains_city_name)
    x["has_known_name"] = x["parsed_values"].apply(_contains_known_name)
    x["has_zip_pattern"] = x["parsed_values"].apply(_contains_zip_code_pattern)
    x["has_phone_pattern"] = x["parsed_values"].apply(_contains_phone_pattern)
    return x


def _build_feature_dataframe(df: pd.DataFrame) -> tuple[pd.DataFrame, list[str]]:
    features = df.columns.tolist()
    value_samples = [_column_sample_values(df[col]) for col in df.columns]
    return _build_feature_dataframe_from_samples(features, value_samples), features


def sanitize_xml_tag(column_name: str) -> str:
    sanitized = re.sub(r"\s+", "_", str(column_name).strip())
    sanitized = XML_TAG_INVALID_CHARS_RE.sub("_", sanitized)
    sanitized = re.sub(r"_+", "_", sanitized).strip("_")
    if not sanitized:
        sanitized = "field"
    if not re.match(r"[A-Za-z_]", sanitized[0]):
        sanitized = f"field_{sanitized}"
    return sanitized


def _write_redacted_json_output(*, redacted_df: pd.DataFrame, redacted_path: str, scan_id: int | None) -> None:
    with open(redacted_path, "w", encoding="utf-8") as handle:
        handle.write('{"scan_id": ')
        json.dump(scan_id, handle)
        handle.write(', "results": [')
        first_record = True
        for row in redacted_df.to_dict(orient="records"):
            if not first_record:
                handle.write(", ")
            json.dump(row, handle, ensure_ascii=False)
            first_record = False
        handle.write("]}")


def _write_redacted_xml_output(*, redacted_df: pd.DataFrame, redacted_path: str) -> None:
    root = ET.Element("results")
    for _, row in redacted_df.iterrows():
        item = ET.SubElement(root, "item")
        for col_name, value in row.items():
            sanitized_tag = sanitize_xml_tag(col_name)
            child = ET.SubElement(item, sanitized_tag)
            child.set("original", str(col_name))
            child.text = "" if value is None else str(value)
    ET.ElementTree(root).write(redacted_path, encoding="utf-8", xml_declaration=True)


def _write_redacted_file(
    redacted_df: pd.DataFrame,
    redacted_path: str,
    ext: str,
    *,
    scan_id: int | None = None,
) -> None:
    if ext == ".csv":
        redacted_df.to_csv(redacted_path, index=False)
    elif ext == ".tsv":
        redacted_df.to_csv(redacted_path, index=False, sep="\t")
    elif ext in {".txt", ".log"}:
        redacted_df.to_csv(redacted_path, index=False)
    elif ext in {".xlsx", ".xls"}:
        redacted_df.to_excel(redacted_path, index=False)
    elif ext == ".json":
        _write_redacted_json_output(redacted_df=redacted_df, redacted_path=redacted_path, scan_id=scan_id)
    elif ext == ".xml":
        _write_redacted_xml_output(redacted_df=redacted_df, redacted_path=redacted_path)
    elif ext == ".docx":
        out = Document()
        if list(redacted_df.columns) == ["text"]:
            for text in redacted_df["text"].fillna("").astype(str).tolist():
                out.add_paragraph(text)
        else:
            for _, row in redacted_df.iterrows():
                out.add_paragraph(" | ".join([str(value) for value in row.values]))
        out.save(redacted_path)
    elif ext == ".html":
        if list(redacted_df.columns) != ["text"]:
            html_out = redacted_df.to_html(index=False)
        else:
            text_blob = "\n".join(redacted_df["text"].fillna("").astype(str).tolist())
            html_out = "<pre>" + html_lib.escape(text_blob) + "</pre>"
        with open(redacted_path, "w", encoding="utf-8") as handle:
            handle.write(html_out)
    else:
        redacted_df.to_csv(redacted_path, index=False)


def _predict_pii_columns(df: pd.DataFrame, model=None) -> tuple[list[str], list[dict[str, object]]]:
    x, features = _build_feature_dataframe(df)
    active_model = model or get_scan_model()
    predictions = active_model.predict(
        x[
            [
                "length",
                "num_underscores",
                "num_digits",
                "has_at",
                "has_email_keyword",
                "pct_email_like",
                "pct_phone_like",
                "pct_ssn_like",
                "pct_ip_like",
                "avg_digits_per_val",
                "avg_val_len",
                "has_dob_pattern",
                "has_gender_term",
                "has_street_suffix",
                "has_city_name",
                "has_known_name",
                "has_zip_pattern",
                "has_phone_pattern",
            ]
        ]
    )
    sample_map = {column: _column_sample_values(df[column]) for column in features}
    pii_columns: list[str] = []
    detection_results: list[dict[str, object]] = []
    for column, is_pii in zip(features, predictions):
        if is_pii != 0:
            continue
        pii_columns.append(column)
        detection_result = _build_detection_result(column_name=column, values=sample_map.get(column, []), ml_detected=True)
        if detection_result is not None:
            detection_results.append(asdict(detection_result))
    return pii_columns, detection_results


def _apply_redactions(
    df: pd.DataFrame,
    pii_columns: list[str],
    aggressive: bool,
) -> tuple[pd.DataFrame, int, int, dict[str, int]]:
    redacted_df = df.copy()
    total_redacted = 0
    total_values = 0
    redacted_type_counts: dict[str, int] = {}

    for col in pii_columns:
        redacted_col, redacted_count, col_total, _, column_type_counts = scan_and_redact_column_with_details(
            df[col], col, aggressive=aggressive
        )
        redacted_df[col] = redacted_col
        total_redacted += redacted_count
        total_values += col_total
        for label, count in column_type_counts.items():
            redacted_type_counts[label] = redacted_type_counts.get(label, 0) + count

    return redacted_df, total_redacted, total_values, redacted_type_counts


def _collect_chunked_metadata(
    *,
    source_path: str,
    ext: str,
    limits: ScanLimits,
) -> tuple[list[str], dict[str, list[str]], int, int]:
    iterator = (
        _iter_csv_like_chunks(source_path=source_path, ext=ext, chunk_rows=limits.chunk_rows)
        if ext in {".csv", ".tsv"}
        else _iter_text_chunks(source_path=source_path, chunk_rows=limits.chunk_rows)
    )

    features: list[str] | None = None
    sample_map: dict[str, list[str]] = {}
    total_rows = 0
    total_cells = 0

    for chunk in iterator:
        if chunk.empty:
            continue
        if features is None:
            features = chunk.columns.tolist()
            if len(features) > limits.max_columns:
                _raise_scan_limit("Uploaded file exceeds allowed processing limits")
            sample_map = {column: [] for column in features}
        elif chunk.columns.tolist() != features:
            raise ValueError("Malformed tabular file.")

        chunk_rows = len(chunk)
        total_rows += chunk_rows
        total_cells += chunk_rows * len(features)
        if total_rows > limits.max_rows or total_cells > limits.max_cells:
            _raise_scan_limit("Uploaded file exceeds allowed processing limits")

        for column in features:
            current_samples = sample_map[column]
            if len(current_samples) >= MAX_FEATURE_SAMPLE_VALUES:
                continue
            values = chunk[column].dropna().astype(str).tolist()
            remaining = MAX_FEATURE_SAMPLE_VALUES - len(current_samples)
            current_samples.extend(values[:remaining])

    if not features or total_rows == 0:
        raise ValueError("File is empty or not supported.")
    return features, sample_map, total_rows, total_cells


def _predict_pii_columns_from_sample_map(
    *,
    features: list[str],
    sample_map: dict[str, list[str]],
    model=None,
) -> tuple[list[str], list[dict[str, object]]]:
    feature_frame = _build_feature_dataframe_from_samples(features, [sample_map.get(column, []) for column in features])
    active_model = model or get_scan_model()
    predictions = active_model.predict(
        feature_frame[
            [
                "length",
                "num_underscores",
                "num_digits",
                "has_at",
                "has_email_keyword",
                "pct_email_like",
                "pct_phone_like",
                "pct_ssn_like",
                "pct_ip_like",
                "avg_digits_per_val",
                "avg_val_len",
                "has_dob_pattern",
                "has_gender_term",
                "has_street_suffix",
                "has_city_name",
                "has_known_name",
                "has_zip_pattern",
                "has_phone_pattern",
            ]
        ]
    )
    pii_columns: list[str] = []
    detection_results: list[dict[str, object]] = []
    for column, is_pii in zip(features, predictions):
        if is_pii != 0:
            continue
        pii_columns.append(column)
        detection_result = _build_detection_result(column_name=column, values=sample_map.get(column, []), ml_detected=True)
        if detection_result is not None:
            detection_results.append(asdict(detection_result))
    return pii_columns, detection_results


def _write_chunked_redacted_output(
    *,
    source_path: str,
    ext: str,
    redacted_path: str,
    pii_columns: list[str],
    limits: ScanLimits,
    aggressive: bool,
) -> tuple[int, int, dict[str, int]]:
    iterator = (
        _iter_csv_like_chunks(source_path=source_path, ext=ext, chunk_rows=limits.chunk_rows)
        if ext in {".csv", ".tsv"}
        else _iter_text_chunks(source_path=source_path, chunk_rows=limits.chunk_rows)
    )
    total_redacted = 0
    total_values = 0
    redacted_type_counts: dict[str, int] = {}
    first_chunk = True

    for chunk in iterator:
        redacted_chunk, redacted_count, chunk_total_values, chunk_type_counts = _apply_redactions(
            chunk,
            pii_columns,
            aggressive,
        )
        total_redacted += redacted_count
        total_values += chunk_total_values
        for label, count in chunk_type_counts.items():
            redacted_type_counts[label] = redacted_type_counts.get(label, 0) + count

        redacted_chunk.to_csv(
            redacted_path,
            index=False,
            sep="\t" if ext == ".tsv" else ",",
            mode="w" if first_chunk else "a",
            header=first_chunk,
        )
        first_chunk = False

    return total_redacted, total_values, redacted_type_counts


def _persist_scan_result(
    *,
    context: ScanContext,
    filename: str,
    ext: str,
    risk_score: int,
    pii_columns: list[str],
    total_redacted: int,
    redacted_type_counts: dict[str, int],
    detection_results: list[dict[str, object]],
    public_redacted_path: str,
) -> int:
    db = None
    try:
        pii_types_string = ",".join(pii_columns) if pii_columns else None
        db = SessionLocal()

        scan = ScanResult(
            user_id=context.user_id,
            company_id=context.company_id,
            filename=filename,
            file_type=ext.lstrip("."),
            risk_score=risk_score,
            pii_types_found=pii_types_string,
            redacted_type_counts=json.dumps(
                build_scan_result_metadata(
                    redacted_type_counts=redacted_type_counts,
                    detection_results=detection_results,
                )
            )
            if redacted_type_counts or detection_results
            else None,
            total_pii_found=total_redacted,
            redacted_file_path=public_redacted_path,
            status="active",
            retention_expiration=build_retention_expiration(
                scanned_at=None,
                retention_days=resolve_retention_days(db, context.company_id, context.tier),
            ),
        )

        db.add(scan)
        db.commit()
        db.refresh(scan)
        logger.info(
            "Saved scan result id=%s user_id=%s file=%s",
            scan.id,
            context.user_id,
            filename,
        )
        record_audit_event(
            db,
            company_id=context.company_id,
            user_id=context.user_id,
            event_type="scan_created",
            event_category="scan",
            description=f"Scan created for file {filename}.",
            target_type="scan",
            target_id=str(scan.id),
            event_metadata={
                "scan_id": scan.id,
                "filename": filename,
                "file_type": ext.lstrip("."),
                "risk_score": risk_score,
                "total_redacted": total_redacted,
            },
        )
        db.commit()
        return scan.id
    except IntegrityError as exc:
        if db is not None:
            db.rollback()
        logger.exception(
            "Integrity error while persisting scan result user_id=%s file=%s",
            context.user_id,
            filename,
        )
        raise RuntimeError(
            "Failed to persist scan result due to integrity constraints. "
            "Verify user/company ids."
        ) from exc
    except Exception as exc:
        if db is not None:
            db.rollback()
        logger.exception(
            "Failed to persist scan result user_id=%s file=%s",
            context.user_id,
            filename,
        )
        raise RuntimeError("Failed to persist scan result.") from exc
    finally:
        if db is not None:
            db.close()


def run_scan_pipeline(
    *,
    file_bytes: bytes | None = None,
    source_path: str | None = None,
    filename: str,
    context: ScanContext,
    model=None,
    limits: ScanLimits | None = None,
    aggressive: bool = False,
) -> ScanPipelineResult:
    ext = os.path.splitext(filename)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError("Unsupported file type.")

    os.makedirs("redacted", exist_ok=True)
    effective_limits = limits or ScanLimits()

    file_id = str(uuid.uuid4())
    redacted_path = os.path.join("redacted", f"redacted_{file_id}{ext}")
    if source_path and ext in {".csv", ".tsv", ".txt", ".log"}:
        features, sample_map, _, _ = _collect_chunked_metadata(
            source_path=source_path,
            ext=ext,
            limits=effective_limits,
        )
        pii_columns, detection_results = _predict_pii_columns_from_sample_map(
            features=features,
            sample_map=sample_map,
            model=model,
        )
        total_redacted, total_values, redacted_type_counts = _write_chunked_redacted_output(
            source_path=source_path,
            ext=ext,
            redacted_path=redacted_path,
            pii_columns=pii_columns,
            limits=effective_limits,
            aggressive=aggressive,
        )
    else:
        df = _parse_to_dataframe(
            file_bytes=file_bytes,
            source_path=source_path,
            filename=filename,
            ext=ext,
            limits=effective_limits,
        )
        pii_columns, detection_results = _predict_pii_columns(df, model=model)
        redacted_df, total_redacted, total_values, redacted_type_counts = _apply_redactions(
            df,
            pii_columns,
            aggressive,
        )

    risk_score = min(round((total_redacted / total_values) * 100), 100) if total_values > 0 else 0

    scan_id = _persist_scan_result(
        context=context,
        filename=filename,
        ext=ext,
        risk_score=risk_score,
        pii_columns=pii_columns,
        total_redacted=total_redacted,
        redacted_type_counts=redacted_type_counts,
        detection_results=detection_results,
        public_redacted_path=redacted_path,
    )

    if not (source_path and ext in {".csv", ".tsv", ".txt", ".log"}):
        _write_redacted_file(redacted_df=redacted_df, redacted_path=redacted_path, ext=ext, scan_id=scan_id)

    logger.info(
        "Completed scan id=%s user_id=%s file=%s redacted_total=%s",
        scan_id,
        context.user_id,
        filename,
        total_redacted,
    )

    return ScanPipelineResult(
        filename=filename,
        pii_columns=pii_columns,
        redacted_file=redacted_path,
        risk_score=risk_score,
        redacted_count=total_redacted,
        total_values=total_values,
        redacted_type_counts=redacted_type_counts,
        detection_results=detection_results,
        scan_id=scan_id,
    )
